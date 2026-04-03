"""
Document text extraction for PDF, DOCX, and image files.

- PDF:   PyMuPDF (fitz) with pymupdf4llm for structured markdown output,
         plus Tesseract OCR fallback for scanned PDFs.
- DOCX:  python-docx with heading/table preservation.
- Image: Tesseract OCR via pytesseract with Pillow preprocessing.
"""

from __future__ import annotations

import io
import logging
import os
import tempfile

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter

logger = logging.getLogger(__name__)

# Supported file types
SUPPORTED_TYPES = {"pdf", "docx", "png", "jpg", "jpeg", "tiff", "bmp"}


# ── Public router ────────────────────────────────────────────────────────────


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Route to the correct extractor based on file type.
    Returns the extracted plain text.
    """
    ft = file_type.lower().strip().strip(".")

    if ft == "pdf":
        return _extract_pdf(file_bytes)
    elif ft == "docx":
        return _extract_docx(file_bytes)
    elif ft in ("png", "jpg", "jpeg", "tiff", "bmp"):
        return _extract_image(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: '{file_type}'. "
            f"Supported types: {', '.join(sorted(SUPPORTED_TYPES))}"
        )


# ── PDF extraction ───────────────────────────────────────────────────────────


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF.
    1. Try PyMuPDF get_text() for native text PDFs.
    2. If text is too short (likely scanned), fall back to OCR on page images.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str] = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            pages_text.append(text.strip())

        full_text = "\n\n".join(pages_text)
        doc.close()

        # If we got meaningful text, return it
        if len(full_text.strip()) > 50:
            logger.info("PDF extracted via PyMuPDF (%d chars)", len(full_text))
            return full_text

        # Otherwise fall back to OCR
        logger.info("PDF has little native text, falling back to OCR")
        return _ocr_pdf_pages(file_bytes)

    except Exception as exc:
        logger.warning("PyMuPDF failed (%s), trying OCR fallback", exc)
        return _ocr_pdf_pages(file_bytes)


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    """Convert each PDF page to an image and run Tesseract OCR."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    texts: list[str] = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        # Render page at 300 DPI for good OCR quality
        pix = page.get_pixmap(dpi=300)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text = _run_ocr(img)
        if text.strip():
            texts.append(text.strip())

    doc.close()
    result = "\n\n".join(texts)
    logger.info("PDF OCR extracted %d chars from %d pages", len(result), len(texts))
    return result


# ── DOCX extraction ──────────────────────────────────────────────────────────


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from a .docx file.
    Preserves headings as markdown-style headers and extracts table content.
    """
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraphs (with heading detection)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            # Extract heading level (Heading 1 → #, Heading 2 → ##, etc.)
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    # Tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append("\n".join(table_rows))

    result = "\n\n".join(parts)
    logger.info("DOCX extracted %d chars", len(result))
    return result


# ── Image extraction (OCR) ───────────────────────────────────────────────────


def _extract_image(file_bytes: bytes) -> str:
    """Extract text from an image using Tesseract OCR with preprocessing."""
    img = Image.open(io.BytesIO(file_bytes))
    text = _run_ocr(img)
    logger.info("Image OCR extracted %d chars", len(text))
    return text


def _run_ocr(image: Image.Image) -> str:
    """
    Preprocess an image and run Tesseract OCR.
    Steps: convert to grayscale → denoise → enhance contrast → resize small images.
    """
    # Convert to grayscale
    img = image.convert("L")

    # Denoise
    img = img.filter(ImageFilter.MedianFilter(size=3))

    # Enhance contrast
    enhancer = ImageEnhance.Contrast(img)
    img = enhancer.enhance(2.0)

    # Up-scale small images for better OCR
    if img.width < 1000:
        scale = 2
        img = img.resize(
            (img.width * scale, img.height * scale),
            Image.Resampling.LANCZOS,
        )

    # Run Tesseract (PSM 6 = assume a single uniform block of text)
    text = pytesseract.image_to_string(img, config="--psm 6")
    return text.strip()
