"""
Unit tests for the text extraction module.
"""

import io
import pytest
from unittest.mock import patch, MagicMock

from app.extractor import extract_text, SUPPORTED_TYPES


class TestExtractTextRouter:
    """Test the file-type routing logic."""

    def test_supported_types_constant(self):
        assert "pdf" in SUPPORTED_TYPES
        assert "docx" in SUPPORTED_TYPES
        assert "png" in SUPPORTED_TYPES
        assert "jpg" in SUPPORTED_TYPES
        assert "jpeg" in SUPPORTED_TYPES
        assert "tiff" in SUPPORTED_TYPES

    def test_unsupported_type_raises(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"dummy", "xyz")

    def test_type_normalization(self):
        """File types should be case-insensitive and handle leading dots."""
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text(b"dummy", ".XYZ")

    def test_strips_dot_from_type(self):
        """extract_text should handle '.pdf' the same as 'pdf'."""
        # This will fail because the bytes aren't a real PDF,
        # but it should NOT fail with "Unsupported file type"
        with pytest.raises(Exception) as exc_info:
            extract_text(b"not-a-pdf", ".PDF")
        assert "Unsupported" not in str(exc_info.value)


class TestDocxExtractor:
    """Test DOCX extraction with a real in-memory .docx file."""

    def test_extracts_paragraphs(self):
        """Create a minimal DOCX in memory and verify extraction."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("This is a test document.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        text = extract_text(buf.read(), "docx")
        assert "Hello World" in text
        assert "test document" in text

    def test_extracts_tables(self):
        """Verify table text is captured."""
        from docx import Document

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "100"

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        text = extract_text(buf.read(), "docx")
        assert "Name" in text
        assert "Alpha" in text
        assert "100" in text

    def test_empty_docx(self):
        """An empty DOCX should return empty string without crashing."""
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        text = extract_text(buf.read(), "docx")
        assert isinstance(text, str)
