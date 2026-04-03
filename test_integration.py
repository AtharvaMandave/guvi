import os
import sys
import base64
import json
from io import BytesIO
from docx import Document

# Load env variables
from dotenv import load_dotenv
load_dotenv(".env")

# Ensure API_KEY is set in environment for the app to pick it up
api_key = os.getenv("API_KEY")
if not api_key or api_key == "your-secret-api-key-here":
    print("WARNING: API_KEY in .env looks like the default.")

from fastapi.testclient import TestClient
from app.main import app

def main():
    print("Creating sample DOCX document...")
    doc = Document()
    doc.add_paragraph("Google to acquire TechStart Inc for $1.5 billion.")
    doc.add_paragraph("The acquisition will be finalized on October 15, 2026. CEO Sundar Pichai announced the deal today, stating it is a great opportunity for growth. Market reactions have been incredibly positive.")
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    file_bytes = buf.read()
    b64_content = base64.b64encode(file_bytes).decode('utf-8')
    
    print("Initializing TestClient and preloading models...")
    # Initialize TestClient with the lifespan context so models load
    with TestClient(app) as client:
        print("Sending request to /document/analyze...")
        response = client.post(
            "/document/analyze",
            headers={"x-api-key": api_key},
            json={
                "fileName": "tech_news.docx",
                "fileType": "docx",
                "fileBase64": b64_content
            }
        )
        
        print(f"\nStatus Code: {response.status_code}")
        if response.status_code == 200:
            print("\nResponse matches structure!")
            result = response.json()
            print(json.dumps(result, indent=2))
        else:
            print("\nError:")
            print(response.text)

if __name__ == "__main__":
    main()
